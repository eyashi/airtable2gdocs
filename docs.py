'''
This will do use the Python docx module until I'm able to actually use the
Google Docs API.
'''
import os
from docx import Document

class WordDoc():
    def __init__(self):
        self.doc = Document()

    def createTestDoc(self):
        self.doc.add_heading('Good Friends')
        self.doc.add_paragraph("What's good friend?")
        self.doc.save('test1.docx')

    def writeDoc(self, entry):
        '''
        Creates a Word document using column headers as paragraph headers
        and the entries as the paragraphs.
        Basic representation of what will eventually form the real document filling
        logic.
        '''
        for field, value in entry['fields'].items():
            if 'Document Title' in field:
                docPath = '%s.docx' % value
                continue

            self.doc.add_heading(field)
            self.doc.add_paragraph(value)

        self.doc.save(docPath)

        return docPath

# this section will be used with document templates... later.
    def openTemplate(self):
        pass
    
    def closeTemplate(self):
        pass

    def writeToTemplate(self):
        pass